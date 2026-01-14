from docx import Document
import os
from openai import OpenAI
import time
import tiktoken
import nltk
from nltk.tokenize import sent_tokenize
import pdfplumber
from collections import Counter
import hashlib
import re
import fitz  # PyMuPDF
from pypdf import PdfReader

def read_docx(path):
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return paragraphs



def is_garbage_line(line):
    line = line.strip()
    if not line:
        return True

    letters = sum(c.isalpha() for c in line)
    non_letters = sum(not c.isalpha() for c in line)

    # too short to be meaningful
    if letters < 3:
        return True

    # mostly punctuation / symbols
    if non_letters > letters:
        return True

    # ornament / page number
    if re.fullmatch(r"[0-9\W_]+", line):
        return True

    return False

def normalize_paragraph(p):
    p = p.lower()
    p = re.sub(r"\s+", " ", p)
    return p.strip()

def read_pdf(path):
    text = ""
    seen_paragraphs = set()

    doc = fitz.open(path)

    for page in doc:
        page_text = page.get_text("text")
        if not page_text:
            continue

        paragraph = ""

        for line in page_text.split("\n"):
            if is_garbage_line(line):
                continue

            if line.strip() == "":
                if paragraph:
                    key = normalize_paragraph(paragraph)
                    if key not in seen_paragraphs:
                        text += paragraph + "\n\n"
                        seen_paragraphs.add(key)
                    paragraph = ""
            else:
                paragraph += (" " if paragraph else "") + line.strip()

        if paragraph:
            key = normalize_paragraph(paragraph)
            if key not in seen_paragraphs:
                text += paragraph + "\n\n"
                seen_paragraphs.add(key)

    sentences = sent_tokenize(text, language="spanish")

    from collections import Counter

    dupes = [s for s, c in Counter(sentences).items() if c > 1]
    print(len(dupes))


    MAX_CHARS = 2000
    chunks = []
    current_chunk = ""

    for sent in sentences:
        if len(current_chunk) + len(sent) + 1 <= MAX_CHARS:
            current_chunk += sent + " "
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sent + " "

    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks


def chunk_paragraphs(paragraphs, max_tokens=1500, model="gpt-4o-mini"):
    enc = tiktoken.encoding_for_model(model)
    chunks = []
    current_chunk = []
    current_tokens = 0

    for para in paragraphs:
        para_tokens = len(enc.encode(para))

        if current_tokens + para_tokens > max_tokens:
            chunks.append("\n\n".join(current_chunk))
            current_chunk = [para]
            current_tokens = para_tokens
        else:
            current_chunk.append(para)
            current_tokens += para_tokens

    if current_chunk:
        chunks.append("\n\n".join(current_chunk))

    return chunks


def translate_chunk(text):
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a professional literary translator. "
                    "Translate the following text from Spanish to English, "
                    "preserving tone, style, and paragraph structure. "
                    "Do not summarize or omit any content."
                )
            },
            {
                "role": "user",
                "content": text
            }
        ],
        temperature=0.2
    )

    return response.choices[0].message.content.strip()


def write_docx(chunks, path):
    doc = Document()
    for chunk in chunks:
        for para in chunk.split("\n\n"):
            doc.add_paragraph(para)
    doc.save(path)

def translate_book(input_file, output_docx,SLEEP_SECONDS):
    if input_file.endswith('.docx'):
        paragraphs = read_docx(input_file)
        chunks = chunk_paragraphs(paragraphs)
    elif input_file.endswith('.pdf'):
        chunks = read_pdf(input_file)
    else:
        print(input_file, ' is not a valid input file')
        return 0

    translated_chunks = []

    for i, chunk in enumerate(chunks, 1):
        print(f"Translating chunk {i}/{len(chunks)}...")
        translated_text = translate_chunk(chunk)
        translated_chunks.append(translated_text)
        time.sleep(SLEEP_SECONDS)

    write_docx(translated_chunks, output_docx)

if __name__ == "__main__":

    SLEEP_SECONDS = 1.2  # avoid rate limits`

    #Picking up chatGPT account info
    api_key = os.getenv("OPENAI_API_KEY")
    client = OpenAI(api_key=api_key)

    translate_book(
        input_file=str($path_to_input_file),
        output_docx=str($path_to_output_file),
        SLEEP_SECONDS
    )